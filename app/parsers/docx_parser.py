"""
DOCX parser dùng python-docx.
Trích xuất text từ paragraphs và tables, sau đó gom thành các page logic.
"""

import logging
from typing import List, Dict

from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

logger = logging.getLogger("ai-service")

DOCX_LOGICAL_PAGE_WORDS = 600


def _iter_document_blocks(doc):
    """
    Yield paragraph/table theo đúng thứ tự trong body DOCX.
    python-docx tách doc.paragraphs và doc.tables riêng nên nếu dùng trực tiếp
    sẽ mất thứ tự table nằm xen giữa paragraph.
    """
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield Table(child, doc)


def _table_to_text(table: Table) -> str:
    rows: list[str] = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
        if cells:
            rows.append(" | ".join(cells))
    return "\n".join(rows).strip()


def _split_long_text(text: str, max_words: int) -> list[str]:
    words = text.split()
    if len(words) <= max_words:
        return [text]
    return [
        " ".join(words[i:i + max_words])
        for i in range(0, len(words), max_words)
    ]


def _build_logical_pages(blocks: list[dict]) -> List[Dict]:
    pages: List[Dict] = []
    current_blocks: list[str] = []
    current_words = 0

    def flush_current() -> None:
        nonlocal current_blocks, current_words
        text = "\n".join(current_blocks).strip()
        if text:
            pages.append(
                {
                    "page_number": len(pages) + 1,
                    "text": text,
                }
            )
        current_blocks = []
        current_words = 0

    for block in blocks:
        text = block["text"]
        is_heading = block.get("is_heading", False)

        if is_heading and current_blocks:
            flush_current()

        for part in _split_long_text(text, DOCX_LOGICAL_PAGE_WORDS):
            part_words = len(part.split())
            if (
                current_blocks
                and current_words + part_words > DOCX_LOGICAL_PAGE_WORDS
            ):
                flush_current()

            current_blocks.append(part)
            current_words += part_words

    flush_current()
    return pages


def parse_docx(file_path: str) -> List[Dict]:
    """
    Parse DOCX file, trả về list page logic.
    DOCX không có page thật ổn định như PDF, nên parser gom nội dung theo
    heading và giới hạn số từ để citation/chunking dễ đọc hơn.

    Returns:
        [{"page_number": 1, "text": "..."}, ...]

    Raises:
        ValueError: Nếu file rỗng.
        RuntimeError: Nếu file hỏng.
    """
    try:
        doc = Document(file_path)
    except Exception as e:
        raise RuntimeError(f"Cannot open DOCX file: {file_path}. Error: {e}")

    blocks: list[dict] = []

    for block in _iter_document_blocks(doc):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if not text:
                continue

            style_name = block.style.name if block.style else ""
            blocks.append(
                {
                    "text": text,
                    "is_heading": style_name.startswith("Heading"),
                }
            )
        elif isinstance(block, Table):
            text = _table_to_text(block)
            if text:
                blocks.append(
                    {
                        "text": text,
                        "is_heading": False,
                    }
                )

    if not blocks:
        raise ValueError(f"No text content extracted from DOCX: {file_path}")

    pages = _build_logical_pages(blocks)

    logger.info(
        "DOCX parsed: %s blocks into %s logical pages from %s",
        len(blocks),
        len(pages),
        file_path,
    )
    return pages
